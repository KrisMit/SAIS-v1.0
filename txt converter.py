"""
DOCUMENT CONVERTER
Extracts text from .docx, .pdf, and other formats
Saves as .txt files ready for the mission agent
"""

import os
import sys
from pathlib import Path

# Try to import document readers
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠ python-docx not installed: pip install python-docx")

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("⚠ PyPDF2 not installed: pip install PyPDF2")

# ─────────────────────────────────────────────────────────────
# EXTRACT TEXT FROM DOCX
# ─────────────────────────────────────────────────────────────
def extract_docx(filepath):
    """Extract text from .docx file"""
    if not HAS_DOCX:
        return None
    try:
        doc = DocxDocument(filepath)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    text.append(row_text)
        return '\n'.join(text)
    except Exception as e:
        print(f"  Error reading {filepath}: {e}")
        return None

# ─────────────────────────────────────────────────────────────
# EXTRACT TEXT FROM PDF
# ─────────────────────────────────────────────────────────────
def extract_pdf(filepath):
    """Extract text from .pdf file"""
    if not HAS_PDF:
        return None
    try:
        text = []
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        print(f"  Error reading {filepath}: {e}")
        return None

# ─────────────────────────────────────────────────────────────
# MAIN CONVERSION
# ─────────────────────────────────────────────────────────────
def convert_all_docs(source_dir, output_dir="mission_docs"):
    """Find and convert all documents"""
    
    source_path = Path(source_dir)
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    
    if not source_path.exists():
        print(f"⚠ Folder not found: {source_dir}")
        return
    
    print(f"\nScanning: {source_dir}")
    print(f"Output folder: {output_dir}/\n")
    
    converted = 0
    
    # DOCX files
    for docx_file in source_path.glob("*.docx"):
        print(f"Converting: {docx_file.name}")
        text = extract_docx(docx_file)
        if text:
            output_file = output_path / f"{docx_file.stem}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"  ✓ Saved: {output_file.name} ({len(text)} chars)")
            converted += 1
        else:
            print(f"  ⚠ Could not extract text")
    
    # PDF files
    for pdf_file in source_path.glob("*.pdf"):
        print(f"Converting: {pdf_file.name}")
        text = extract_pdf(pdf_file)
        if text:
            output_file = output_path / f"{pdf_file.stem}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"  ✓ Saved: {output_file.name} ({len(text)} chars)")
            converted += 1
        else:
            print(f"  ⚠ Could not extract text")
    
    # TXT files (copy them)
    for txt_file in source_path.glob("*.txt"):
        print(f"Copying: {txt_file.name}")
        with open(txt_file, 'r', encoding='utf-8') as f:
            text = f.read()
        output_file = output_path / txt_file.name
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"  ✓ Copied: {output_file.name}")
        converted += 1
    
    print(f"\n{'='*56}")
    print(f"✓ Converted {converted} documents to {output_dir}/")
    print(f"{'='*56}\n")
    print(f"Now run: python mission_agent_local.py")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        source = sys.argv[1]
    else:
        # Default: look in same folder as this script
        source = input("Enter source folder path (where your docs are): ").strip()
    
    if not source:
        print("No folder specified")
        sys.exit(1)
    
    convert_all_docs(source)